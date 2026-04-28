import io
import os
import re
import zipfile
import tempfile
import subprocess
from datetime import datetime
from typing import Dict, Set

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt
from pydantic import BaseModel, Field, validator

st.set_page_config(page_title="Gerador de MOU", page_icon="📝", layout="wide")

PLACEHOLDER_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}", re.IGNORECASE)

# ---------------------------
# DOCX utils
# ---------------------------
def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p


def _para_text(p) -> str:
    txt = "".join(run.text for run in p.runs) or p.text or ""
    return re.sub(r"\s+", " ", txt).strip()


def extract_placeholders(doc: Document) -> Set[str]:
    found = set()
    for p in _iter_all_paragraphs(doc):
        for m in PLACEHOLDER_RE.finditer(_para_text(p)):
            found.add(m.group(1).strip().upper())
    return found


# ---------------------------
# Exceções de negrito
# ---------------------------
def is_exception(text: str) -> bool:
    t = text.lower().strip()

    if "{{bp_date}}" in t or "{{comments}}" in t or "{{comments_eng}}" in t:
        return True

    if "como parte integrante deste documento" in t:
        return True

    if "as an integral part of this document" in t:
        return True

    if "business plan" in t:
        return True

    if "arquivo:" in t or "file:" in t:
        return True

    if "especifica" in t or "specification" in t:
        return True

    if re.fullmatch(r"n\s*/?\s*a\.?", t):
        return True

    return False


# ---------------------------
# Replace placeholders
# ---------------------------
def replace_doc(doc: Document, mapping: Dict[str, str]):
    exceptions = set()

    normalized_mapping = {
        k.strip().strip("{} ").upper(): "" if pd.isna(v) else str(v)
        for k, v in mapping.items()
    }

    for p in _iter_all_paragraphs(doc):
        original = _para_text(p)

        if is_exception(original):
            exceptions.add(p)

        replaced = original
        for k, v in normalized_mapping.items():
            pattern = re.compile(r"\{\{" + re.escape(k) + r"\}\}", re.IGNORECASE)
            replaced = pattern.sub(v, replaced)

        if replaced != original:
            for _ in range(len(p.runs)):
                p._element.remove(p.runs[0]._element)
            p.add_run(replaced)

    return exceptions


# ---------------------------
# Formatação
# ---------------------------
def format_doc(doc: Document, exceptions: Set):
    paragraphs = list(_iter_all_paragraphs(doc))

    # Calibri 11 + negrito em tudo
    for p in paragraphs:
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(11)
            r.bold = True

    # Remove negrito das exceções
    for i, p in enumerate(paragraphs):
        text = _para_text(p).lower().strip()

        if p in exceptions or is_exception(text):
            for r in p.runs:
                r.bold = False

        # Caso o "2." esteja separado da frase
        if text in ["2.", "2"]:
            for r in p.runs:
                r.bold = False

            if i + 1 < len(paragraphs):
                next_text = _para_text(paragraphs[i + 1]).lower()
                if is_exception(next_text):
                    for r in paragraphs[i + 1].runs:
                        r.bold = False


# ---------------------------
# PDF
# ---------------------------
def convert_pdf(docx_bytes: bytes):
    with tempfile.TemporaryDirectory() as td:
        docx_path = os.path.join(td, "file.docx")
        pdf_path = os.path.join(td, "file.pdf")

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            with open(pdf_path, "rb") as f:
                return f.read()
        except Exception:
            pass

        try:
            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    td,
                    docx_path,
                ],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            with open(pdf_path, "rb") as f:
                return f.read()
        except Exception:
            return None


# ---------------------------
# Validação
# ---------------------------
class JobConfig(BaseModel):
    title: str
    placeholders: Dict[str, str] = Field(default_factory=dict)

    @validator("placeholders")
    def normalize_keys(cls, v: Dict[str, str]):
        fixed = {}
        for k, val in v.items():
            key = k.strip().strip("{} ").upper()
            fixed[key] = "" if pd.isna(val) else str(val)
        return fixed


# ---------------------------
# UI
# ---------------------------
st.title("Gerador de MOU")
st.caption("Upload do template DOCX, preenchimento manual ou via XLSX, e download em DOCX/PDF.")

with st.sidebar:
    st.header("⚙️ Configuração")
    template_file = st.file_uploader("Upload template DOCX", type=["docx"])
    batch_mode = st.toggle("Modo Excel (.xlsx) em lote", value=False)

if not template_file:
    st.info("Envie o template DOCX para começar.")
    st.stop()

template_bytes = template_file.read()
doc_template = Document(io.BytesIO(template_bytes))
fields = sorted(list(extract_placeholders(doc_template)))

if not fields:
    st.warning("Nenhum placeholder encontrado. Use o formato {{CHAVE}} no template.")
    st.stop()

# ---------------------------
# Modo individual
# ---------------------------
if not batch_mode:
    st.subheader("Gerar 1 documento")

    inputs = {}
    cols = st.columns(3)

    for i, field in enumerate(fields):
        with cols[i % 3]:
            inputs[field] = st.text_input(field)

    default_title = f"MOU – {inputs.get('FANTASY_NAME', 'Sem Nome')} – {datetime.now().strftime('%Y-%m-%d')}"
    title = st.text_input("Nome do arquivo", value=default_title)

    if st.button("Gerar documento", type="primary"):
        cfg = JobConfig(title=title, placeholders=inputs)

        doc = Document(io.BytesIO(template_bytes))
        exceptions = replace_doc(doc, cfg.placeholders)
        format_doc(doc, exceptions)

        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()

        st.success("Documento gerado!")

        col1, col2 = st.columns(2)

        with col1:
            st.download_button(
                "⬇️ Baixar DOCX",
                data=docx_bytes,
                file_name=f"{cfg.title}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        with col2:
            pdf_bytes = convert_pdf(docx_bytes)
            if pdf_bytes:
                st.download_button(
                    "⬇️ Baixar PDF",
                    data=pdf_bytes,
                    file_name=f"{cfg.title}.pdf",
                    mime="application/pdf",
                )
            else:
                st.info("PDF não disponível neste ambiente. O DOCX foi gerado normalmente.")


# ---------------------------
# Modo XLSX em lote
# ---------------------------
else:
    st.subheader("Gerar documentos em lote via Excel")

    st.markdown(
        "A planilha `.xlsx` deve ter colunas com os mesmos nomes dos placeholders, sem `{{}}`. "
        "Ex.: `FANTASY_NAME`, `GROUP_NAME`, `CNPJ`, `RESPONSIBLE_NAME`, `COMMENTS`, `COMMENTS_ENG`. "
        "Opcional: coluna `TITLE` para nomear cada arquivo."
    )

    xlsx_file = st.file_uploader("Upload Excel (.xlsx)", type=["xlsx"])

    if xlsx_file and st.button("Gerar lote", type="primary"):
        df = pd.read_excel(xlsx_file, engine="openpyxl")

        missing_cols = [f for f in fields if f not in df.columns]
        if missing_cols:
            st.warning("Colunas ausentes no XLSX: " + ", ".join(missing_cols))

        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as z:
            for idx, row in df.iterrows():
                data = {
                    field: "" if pd.isna(row.get(field, "")) else str(row.get(field, ""))
                    for field in fields
                }

                title = row.get("TITLE", "")
                if pd.isna(title) or str(title).strip() == "":
                    title = f"MOU – {data.get('FANTASY_NAME', data.get('GROUP_NAME', 'Documento'))}"

                title = str(title).strip()

                doc = Document(io.BytesIO(template_bytes))
                exceptions = replace_doc(doc, data)
                format_doc(doc, exceptions)

                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                docx_bytes = docx_buffer.getvalue()

                z.writestr(f"{title}.docx", docx_bytes)

                pdf_bytes = convert_pdf(docx_bytes)
                if pdf_bytes:
                    z.writestr(f"{title}.pdf", pdf_bytes)

        zip_buffer.seek(0)

        st.success("Lote gerado!")
        st.download_button(
            "⬇️ Baixar ZIP",
            data=zip_buffer,
            file_name="mous_gerados.zip",
            mime="application/zip",
        )


with st.expander("Campos encontrados no template"):
    st.write(fields)
